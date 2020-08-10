"""
python3 script for Lijuan
Author: Zhang Chengsheng, @2019.07.08
"""

import os,sys,docx,time
from docx.shared import Cm
import abi_new
import shutil


def config_read(config,png_dir):
    id_dict = {}
    png_dict = {}
    header_dict = {}
    abi_dict = {}
    abi_list = [i for i in os.listdir(abi_dir) if i.endswith('.ab1')]
    png_list = [i for i in os.listdir(png_dir) if i.endswith('.mdf.png')]
    header_list = [i for i in os.listdir(png_dir) if i.endswith('.header.png')]
    with open(config,'r') as f:
        for i in f.readlines():
            if i.startswith('#') or not i or i == '\n':
                continue
            line = i.strip('\n').split('\t')
            id = line[0]
            if id in id_dict:
                id_dict[id].append(line)
            else:
                id_dict[id] = [line]

            # add png_dict
            pcr = line[5]
            name_base = id.split('/')
            prefix = name_base[0][:-2] if name_base[0][-1].isdigit() else name_base[0][:-1]
            suffix = []
            for i in id.lstrip(prefix).split('/'):
                suffix.append(i)
            #names = [prefix+i for i in suffix]
            pngs = [pcr+'_'+i for i in suffix]
            for j in header_list:
                if pcr in j:
                    header_dict[pcr] = os.path.join(png_dir,j)
                    break
            for idx,j in enumerate(pngs):
                for k in abi_list:
                    if j in k:
                        abi_dict[j] = os.path.join(abi_dir,k)
                        break
            for idx,j in enumerate(pngs):
                for k in png_list:
                    if pngs[idx] in k:
                        png_path = k
                        break
                else:
                    png_path = 0

                png_dict[j] = os.path.join(png_dir,png_path) if png_path else ''
    #for i in abi_dict:
    #    print(i,abi_dict[i])
    return id_dict,png_dict,header_dict,abi_dict


class word_make:
    def __init__(self,id,info,pngs,header,outdir,abi,o):
        self.id = id
        self.info = info
        self.template = r'D:\zcs-genex\SCRIPTS\Scripts_for_work\lijuan_saifu_report\N4770报告template.docx'
        self.outdir = outdir
        self.pngs = pngs
        self.header = header
        self.abi = abi
        self.o = o

    def _get_row_num(self):
        return len(self.info)

    def _makedir(self,path):
        #base = os.path.dirname(path)
        if not os.path.exists(path):
            os.makedirs(path)

    def _get_names(self):
        name = self.info[0][0]
        name_base = name.split('/')
        name_prefix = name_base[0][:-2] if name_base[0][-1].isdigit() else name_base[0][:-1]
        suffix = []
        for i in name.lstrip(name_prefix).split('/'):
            suffix.append(i)
        names = [name_prefix + i for i in suffix]
        date = self.info[0][1].replace('/','-')
        date2 = self.info[0][4].replace('/', '-')
        return names,date,name_prefix,suffix,date2

    def shutil_copy(self,source,target_dir):
        file_name = os.path.basename(source)
        new_file = os.path.join(target_dir,file_name)
        shutil.copy(source,new_file)

    def _set_gene_sheet_width(self,sheet,row_num):
        width = [Cm(2),Cm(2.5),Cm(3.5),Cm(4.5)]
        for col in range(4):
            for row in range(row_num):
                sheet.cell(row,col).width = width[col]

    def _add_gene_sheet(self,gene_info,D):
        D.add_paragraph('\n\n一代验证结果')
        names, date, name_prefix,suffix,dat32 = self._get_names()
        gene = gene_info[2]
        gene_site = gene_info[3]
        pcr = gene_info[5]
        table2_row_num = len(names) + 1
        table2 = D.add_table(rows=table2_row_num, cols=4, style='Table Grid')
        table2_header = ['姓名','基因','位点(hg38)','突变类型']
        self._set_gene_sheet_width(table2,table2_row_num)

        for i in range(4):
            table2.rows[0].cells[i].text = table2_header[i]
            for j in range(table2_row_num-1):
                table2.rows[j+1].cells[0].text = names[j]
                table2.rows[j+1].cells[1].text = gene
                table2.rows[j+1].cells[2].text = gene_site
        D.add_paragraph('\nNCBI 参考序列')
        if pcr in self.header:
            D.add_picture(self.header[pcr],width=Cm(18), height=Cm(0.5))
            ##self.shutil_copy(self.header[pcr],self.outdir)
        else:
            print('{}|{}|{}没有参考序列图片...'.format(name_prefix,gene,pcr))
            self.o.write('\t{}|{}|{}没有参考序列图片...\n'.format(name_prefix,gene,pcr))
        for idx,i in enumerate(names):
            D.add_paragraph('\n'+i+'测序序列')
            png = pcr + '_' + suffix[idx]
            if png in self.pngs:
                if self.pngs[png]:
                    D.add_picture(self.pngs[png],width=Cm(18), height=Cm(5))
                    genotypes = self.pngs[png].split('.')[-3]
                    ref = genotypes.split('-')[-1]
                    seq = genotypes.split('-')[0]
                    if seq[0] == seq[1]:
                        if seq[0] != ref:
                            text = '纯合突变{}/{}'.format(seq[0],seq[0])
                        else:
                            text = '野生型{}/{}'.format(seq[0], seq[0])
                    else:
                        if seq[0] != ref:
                            text = '杂合突变{}/{}'.format(seq[0],seq[1])
                        else:
                            text = '杂合突变{}/{}'.format(seq[1],seq[0])
                    table2.rows[idx+1].cells[3].text = text
                    ##self.shutil_copy(self.pngs[png], self.outdir)
                    self.shutil_copy(self.abi[png], self.outdir)
                else:
                    print('{}|{}|{} 没有峰图文件...'.format(i, gene, pcr))
                    self.o.write('\t{}|{}|{} 没有峰图文件...\n'.format(i, gene, pcr))
            else:
                print('{}|{}|{} 没有峰图文件...'.format(i,gene,pcr))
                self.o.write('\t{}|{}|{} 没有峰图文件...\n'.format(i, gene, pcr))

    def word_table_cell_insert_pic(table, cell, pic):
        run = table.cell(cell[0], cell[1]).paragraphs[0].add_run()
        run.add_picture(pic, width=Cm(18), height=Cm(3))

    def run(self):
        D = docx.Document(self.template)
        names,date,name_prefix,suffix,date2 = self._get_names()
        date3 = time.strftime("%Y-%m-%d", time.localtime())
        table1_row_num = len(names) + 1
        table1 = D.add_table(rows=table1_row_num,cols=7,style='Table Grid')
        table1_header = ['姓名','样本类型','研究项目','下单时间','到样时间','报告时间','检测目的']
        table1_cell = [date,'血液','赛福Sanger','',date2,date3,'一代验证']

        self.outdir = os.path.join(self.outdir, date)
        self.outdir = os.path.join(self.outdir, name_prefix)
        self._makedir(self.outdir)

        for i in range(7):
            table1.rows[0].cells[i].text = table1_header[i]
            for j in range(table1_row_num-1):
                table1.rows[j+1].cells[i].text = table1_cell[i]
                table1.rows[j+1].cells[0].text = names[j]
                table1.rows[j+1].cells[3].text = date

        for i in self.info:
            self._add_gene_sheet(i,D)

        self.out = os.path.join(self.outdir,name_prefix+'.docx')
        D.save(self.out)


def main(config,png_dir,out):
    info,pngs,header,abi = config_read(config,png_dir)
    report = os.path.join(out,'log.txt')
    with open(report,'w') as o:
        o.write('报告日期: {}\n\n'.format(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())))
        for i in info:
            #print(i,info[i])
            print('正在处理: {}'.format(i))
            o.write('{}:\n'.format(i))
            A = word_make(i,info[i],pngs,header,out,abi,o)
            A.run()


def debug():
    global abi_dir
    png_dir = r'H:\09-27-赛福\测序结果\pngs'
    abi_dir = r'H:\09-27-赛福\测序结果'
    #png_dir = abi_new.abifile_parse(abi_dir)
    #abi_new.png_reshape_batch(png_dir)
    config = r'H:\09-27-赛福\config.txt'
    out = r'H:\09-27-赛福\report'
    main(config, png_dir, out)
    exit()


def welcome():
    sys.stdout.write('======================================\n')
    sys.stdout.write('===  Script for ?saifu? report     ===\n')
    sys.stdout.write('===  Author: Zhang Chengsheng      ===\n')
    sys.stdout.write('===  @2019.07.12                   ===\n')
    sys.stdout.write('===  Usage: No usage               ===\n')
    sys.stdout.write('======================================\n\n')


if __name__ == '__main__':
    #debug()
    welcome()
    while 1:
        config = input("请输入样本配置文件:\n")
        if not os.path.isfile(config):
            sys.stdout.write("配置文件不存在，")
        else:
            break

    while 1:
        abi_dir = input("请输入abi文件夹路径:\n")
        if not os.path.isdir(abi_dir):
            sys.stdout.write("文件夹路径有误，")
        else:
            break

    while 1:
        out = input("请输入报告输出路径:\n")
        if not os.path.isdir(out):
            sys.stdout.write("文件夹路径有误，")
        else:
            break

    pic_only = input("DEBUG模式:\n")
    if 1:
    #try:
        if pic_only != '1':
            png_dir = abi_new.abifile_parse(abi_dir)
            abi_new.png_reshape_batch(png_dir)
        else:
            png_dir = os.path.join(abi_dir,'pngs')
        main(config, png_dir, out)

    sys.stdout.write('\n')
    input("运行结束，回车退出。\n")
