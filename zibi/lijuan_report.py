"""
python3 script
Author: Zhang Chengsheng, @2018.11.12
Usage: 

2020.08.31
update to 18 sites
"""

import os,sys,docx,time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
import abi_new
from win32com.client import Dispatch,constants,gencache


def word_table_cell_replace(table,cell,txt,warning=False,line=0):
    font = "Songti SC"
    run = table.cell(cell[0],cell[1]).paragraphs[line].add_run(txt)
    run.font.name = font
    run.font.size = Pt(9)
    table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if warning:
        run.font.highlight_color = WD_COLOR_INDEX.RED


def word_table_cell_insert_pic(table,cell,pic):
    run = table.cell(cell[0],cell[1]).paragraphs[0].add_run()
    run.add_picture(pic,width=Cm(6.78),height=Cm(2.8))


def config_parsing(config,pic_dir):
    config_dict = {}
    point_num = 18
    with open(config,'r') as f:
        for i in f.readlines():
            if i.startswith("#") or i == '\n':
                continue
            lines = i.strip().split('\t')
            if len(lines) != 6:
                sys.stdout.write("ERROR: 配置文件格式有误!\n错误行：{}\n".format(i))
                continue
            seq_num = lines[5]
            if seq_num not in config_dict:
                config_dict[seq_num] = lines[:-1]
            else:
                sys.stdout.write("WARNING: 测序编号重复:{}, 后者未能录入!\n".format(seq_num))

    pic_dict = {}
    for i in config_dict:
        pic_dict[i] = []
        for k in range(point_num):
            for file in os.listdir(pic_dir):
                if "%s-%s" %(str(i),str(k+1)) in file or "%s_%s" %(str(i),str(k+1)) in file:
                    if file.endswith('.0'):
                        sys.stdout.write("WARNING: 样本{}-{}的错配超限，对应位点将被留空!\n".format(i, k + 1))
                        pic_dict[i].append("")
                        break
                    pic_dict[i].append(os.path.join(pic_dir, file))
                    break
            else:
                sys.stdout.write("WARNING: 样本{}中没有发现{}-{}的图片文件! 对应位点将被留空!\n".format(i, i, k + 1))
                print("%s-%s" %(str(i),str(k+1)))
                pic_dict[i].append("")
    return config_dict,pic_dict


def word_work(model,out,context,pics):
    site_dict = {1:["无","Val158Met"],
                 2:["无","Arg297="],
                 3:["--","--"],
                 4:["无","Ala222Val"],
                 5:["无","Pro80="],
                 6:["无","Asp919Gly"],
                 7:["无","Ile22Met"],
                 8:["--","--"],
                 9:["--","--"],
                 10:["--","--"],
                 11:["--","--"],
                 12:["无","Tyr233="],
                 13:["无","Ala360="],
                 14:["无","Ser427="],
                 15:["无","Leu435Phe"],
                 16:["无","Asp298Glu"],
                 17:["--","--"],
                 18: ["无", "Thr92Ala"],
                 }
    key = False
    d = docx.Document(model)
    table1 = d.tables[0]
    table3 = d.tables[2]
    table4 = d.tables[3]
    table5 = d.tables[4]
    word_table_cell_replace(table1, (0, 1), context[1])
    word_table_cell_replace(table1, (2, 1), context[0])
    word_table_cell_replace(table1, (2, 3), context[2])
    word_table_cell_replace(table1, (2, 5), context[3])
    word_table_cell_replace(table5, (0, 1), context[4])
    word_table_cell_replace(table5, (0, 3), context[3])
    for idx,i in enumerate(pics):
        if i == "":
            key = idx+1
        else:
            word_table_cell_insert_pic(table4, ((idx//2)+2, 2*(idx%2)+1), i)
            #genotype, res, type, miss_count = seq_resolve.base_identity(i,refseq_dict[idx+1][0],refseq_dict[idx+1][1])
            genotype = os.path.basename(i).split('.')[1].split('-')[0].upper()
            res = os.path.basename(i).split('.')[1].split('-')[1].upper()
            ###
            if genotype[0] == genotype[1]:
                if genotype[0] == res:
                    seq_res = "-/-"
                    type = "野生型"
                else:
                    seq_res = "+/+"
                    type = "纯合"
                target_seq = "{}/{}".format(genotype[0], genotype[0])
            else:
                if genotype[0] == res:
                    target_seq_2 = genotype[0]
                    target_seq_1 = genotype[1]
                    seq_res = "+/-"
                    type = "杂合"
                elif genotype[1] == res:
                    target_seq_2 = genotype[1]
                    target_seq_1 = genotype[0]
                    seq_res = "+/-"
                    type = "杂合"
                else:
                    target_seq_2 = genotype[1]
                    target_seq_1 = genotype[0]
                    seq_res = "+/+"
                    type = "杂合"
                if idx in []:
                #if idx in [2,6]:  # 需要互补的点
                    target_seq_1 = abi_new.trans(target_seq_1)
                    target_seq_2 = abi_new.trans(target_seq_2)
                target_seq = "{}/{}".format(target_seq_1, target_seq_2)


            #warning = True if miss_count > 2 else False
            #if warning:
                #sys.stdout.write("Warning: {} 错配超限！".format(out))
            warning = False
            word_table_cell_replace(table3, (idx + 2, 4), target_seq, warning=warning)
            word_table_cell_replace(table3, (idx + 2, 5), seq_res, warning=warning)
            word_table_cell_replace(table3, (idx + 2, 6), type, warning=warning)
            word_table_cell_replace(table4, ((idx//2)+2, 2*(idx%2)), "({})".format(type),warning=warning,line=-1)

            if type != "野生型":
                word_table_cell_replace(table3, (idx + 2, 7), site_dict[idx + 1][1],warning=warning)
            else:
                word_table_cell_replace(table3, (idx + 2, 7), site_dict[idx + 1][0],warning=warning)
    d.save(out)
    return key


def word2pdf(docx_file,pdf_file):
    time.sleep(1)
    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
    w = Dispatch("Word.Application")
    doc = w.Documents.Open(docx_file, ReadOnly=1)
    doc.ExportAsFixedFormat(pdf_file, constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    #sys.stdout.write("WARNING: docx转pdf发生错误！错误文件:{}\n".format(docx_file))
    w.Quit(constants.wdDoNotSaveChanges)


def main(config,pic_dir,out_dir):
    if not os.path.exists(out_dir):
        os.mkdir(out_dir)
    base_model = r"模板路径"
    cfg, pic_dict = config_parsing(config,pic_dir)
    for i in cfg:
        sys.stdout.write("生成{}报告...编号：{}...\n".format(cfg[i][1],i))
        out_file = os.path.join(out_dir,"基因位点筛查报告单-{}-.docx".format(cfg[i][1]))
        key = word_work(base_model,out_file,cfg[i],pic_dict[i])
        if not key:
            try:
                word2pdf(out_file,out_file.replace(".docx",".pdf"))
            except:
                sys.stdout.write("WARNING: docx转pdf发生错误！错误文件:{}\n".format(out_file))
        else:
            sys.stdout.write("WARNING:{}没有生成PDF文件，编号{}，{}\n".format(cfg[i][1],i,key))


def welcome():
    sys.stdout.write('======================================\n')
    sys.stdout.write('===  Author: Zhang Chengsheng      ===\n')
    sys.stdout.write('===  @2018.11.09                   ===\n')
    sys.stdout.write('===  Usage: No usage               ===\n')
    sys.stdout.write('======================================\n\n')


if __name__ == '__main__':
    welcome()
    while 1:
        config = input("请输入样本配置文件:\n")
        if not os.path.isfile(config):
            sys.stdout.write("配置文件不存在，")
        else:
            break

    while 1:
        pic_dir = input("请输入abi文件夹路径:\n")
        if not os.path.isdir(pic_dir):
            sys.stdout.write("文件夹路径有误，")
        else:
            break

    while 1:
        out = input("请输入报告输出路径:\n")
        if not os.path.isdir(out):
            sys.stdout.write("文件夹路径有误，")
        else:
            break

    pic_only = input("DEBUG模式:\n")

    try:
        if pic_only != '1':
            dir1 = abi_new.abifile_parse(pic_dir)
            abi_new.png_reshape_batch(dir1)
        else:
            dir1 = os.path.join(pic_dir,'pngs')
        main(config, dir1, out)
    except Exception as e:
        print(e)
        sys.stdout.write('\n')
    sys.stdout.write('\n')
    input("运行结束，回车退出\n")


