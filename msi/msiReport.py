"""
Author: Zhang Chengsheng, @2020.05.15
python3 script
娟姐MSI报告用
"""

import os,sys
import docx
import pandas
import openpyxl
from docx.shared import Pt,Cm
import pdfRead


def txtConfigParse(txtInput):
    res = {}
    with open(txtInput,'r') as f:
        for i in f.readlines():
            if 'Size3' in i:
                print('!!!!!!!!!!!! Size3 in files !!!!!!!!!!!')
                exit(10086)
            if not i.strip() or i.startswith('Sample File'):
                continue
            line = i.strip().split('\t')
            sampleName = line[1]
            panel = line[2]
            marker = line[3].replace('-','').lower()
            size1 = line[7]
            size2 = line[8]
            if sampleName in res:
                if marker not in res[sampleName]:
                    res[sampleName][marker] = [panel,size1,size2]
            else:
                res[sampleName] = {marker:[panel,size1,size2]}
    return res


def xlsConfigParse(xlsInput):
    res = {}
    wb = openpyxl.load_workbook(xlsInput)
    sheet = wb[wb.sheetnames[0]]
    for row in sheet.rows:
        if row[0].value in ['患者名称']:
            continue
        Pname = row[0].value if row[0].value else Pname
        Nid = row[2].value.replace('_','-')
        Tid = row[3].value.replace('_','-')
        if Pname in res:
            res[Pname][0].append(Nid)
            res[Pname][1].append(Tid)
        else:
            res[Pname] = [[Nid],[Tid]]
    return res


def figConfigMake(pngDir):
    return {i.strip('.png').split('_')[-1]: os.path.join(pngDir,i) for i in os.listdir(pngDir) if i.endswith('.png')}


class markerCheck:
    def __init__(self,normal,tumor,txtConfig,markerLib):
        self.normal = normal
        self.tumor = tumor
        self.total = txtConfig
        self.markerLib = markerLib
        self.T = TNmarker(self.markerLib)
        self.N = TNmarker(self.markerLib)
        self.markerInitial()
        self.markerReplace()

    def markerInitial(self):
        for i in self.markerLib:
            self.T.i = []
            self.N.i = []

    def markerReplace(self):
        for i in self.normal:
            if i in self.total:
                for j in self.total[i]:
                    if j in self.markerLib:
                        self.N.add(j,self.total[i][j])
                    else:
                        print('Unknown marker N: {}'.format(j))
        for i in self.tumor:
            if i in self.total:
                for j in self.total[i]:
                    if j in self.markerLib:
                        self.T.add(j,self.total[i][j])
                    else:
                        print('Unknown marker T: {}'.format(j))


class TNmarker:
    def __init__(self,markerLib):
        self.value = {}
        for i in markerLib:
            self.__dict__[i] = []

    def add(self,id,value):
        self.__dict__[id] = value
        self.value[id] = value

    def __str__(self):
        return str(self.value)


def report(txtConfig,xlsConfig,figConfig,templet,outputDir):
    markerLib = ['nr21','bat26','pentac','nr27','bat25','pentad','amel','nr24','mono27']
    refLib = ['pentac','pentad','amel']
    if not os.path.exists(outputDir):
        os.makedirs(outputDir)
    log = os.path.join(outputDir,'log.txt')
    o = open(log,'w')
    for name in xlsConfig:
        ## todo:table[0]替换name
        print('正在处理:{}...'.format(name))
        o.write('正在处理:{}...\n'.format(name))
        reportName = os.path.join(outputDir,'{}.docx'.format(name))
        D = docx.Document(templet)
        word_table_cell_replace(D.tables[0],[0,1],name)
        normal = xlsConfig[name][0]
        tumor = xlsConfig[name][1]
        Marker = markerCheck(normal,tumor,txtConfig,markerLib)
        diff = 0
        for idx,i in enumerate(markerLib):
            size1T = Marker.T.value[i][1]
            size2T = Marker.T.value[i][2]
            size1N = Marker.N.value[i][1]
            size2N = Marker.N.value[i][2]
            if i not in refLib:
                if size2N and not size2T:
                    diff += 1
                elif not size2N and size2T:
                    diff += 1
            word_table_cell_replace(D.tables[1], [idx + 2, 1], size1N)
            word_table_cell_replace(D.tables[1], [idx + 2, 2], size2N)
            word_table_cell_replace(D.tables[1], [idx + 2, 4], size1T)
            word_table_cell_replace(D.tables[1], [idx + 2, 5], size2T)

        for idx,i in enumerate(normal):
            if i in figConfig:
                word_table_cell_insert_pic(D.tables[2], [idx, 0], figConfig[i])
            else:
                print('{}没有峰图文件'.format(i))
                o.write('{}没有峰图文件\n'.format(i))
        for idx,i in enumerate(tumor):
            if i in figConfig:
                word_table_cell_insert_pic(D.tables[2], [idx, 1], figConfig[i])
            else:
                print('{}没有峰图文件'.format(i))
                o.write('{}没有峰图文件\n'.format(i))

        if diff == 1:
            msiType = '微卫星低度不稳定(MSI-L)型'
        elif diff > 1:
            msiType = '微卫星高度不稳定(MSI-H)型'
        else:
            msiType = '微卫星稳定(MSS)型'
        run = D.paragraphs[8].add_run(msiType)
        run.font.size = Pt(11)
        print('{}为{}'.format(name,msiType))
        o.write('{}为{}\n'.format(name,msiType))
        D.save(reportName)
    o.close()

def word_table_cell_insert_pic(table, cell, pic):
    run = table.cell(cell[0], cell[1]).paragraphs[0].add_run()
    run.add_picture(pic, width=Cm(8.5), height=Cm(3.33))


def word_table_cell_replace(table,cell,txt,line=0):
    run = table.cell(cell[0],cell[1]).paragraphs[line].add_run(txt)


def wordTest():
    file = r'D:\zcs-genex\SCRIPTS\Scripts_for_work\lijuan_pdf\MSI报告模板V20200515.docx'
    D = docx.Document(file)
    print(D.tables[2].cell(1,0).paragraphs[0].text)


def run(baseDir,debug=''):
    templet = r'D:\zcs-genex\SCRIPTS\Scripts_for_work\lijuan_pdf\MSI报告模板V20200515.docx'
    pngDir = os.path.join(baseDir,'pngs')
    reportDir = os.path.join(baseDir,'reports')
    txt = [i for i in os.listdir(baseDir) if i.endswith('.txt') or i.endswith('.TXT')]
    xls = [i for i in os.listdir(baseDir) if i.endswith('.xls') or i.endswith('.xlsx') or i.endswith('.XLSX') or i.endswith('.XLS')]
    if len(txt) != 1:
        print('存在未知 txt 文件！')
        exit(10000)
    else:
        txt = txt[0]
    if len(xls) != 1:
        print('存在未知 excel 文件！')
        exit(10001)
    else:
        xls = xls[0]
    if debug != '1':
        pdfRead.run(baseDir,pngDir)
    xlsConfig = xlsConfigParse(xls)
    txtConfig = txtConfigParse(txt)
    figConfig = figConfigMake(pngDir)
    report(txtConfig, xlsConfig, figConfig, templet, reportDir)


def welcome():
    sys.stdout.write('======================================\n')
    sys.stdout.write('===  Author: Zhang Chengsheng      ===\n')
    sys.stdout.write('===  @2020.05.15                   ===\n')
    sys.stdout.write('===  for MSI report                ===\n')
    sys.stdout.write('======================================\n\n')


def main():
    welcome()
    while 1:
        baseDir = input("请输入目标路径:\n")
        if not os.path.isdir(baseDir):
            sys.stdout.write("目标路径不存在，")
        else:
            break
    debug = input("DEBUG模式:\n")
    run(baseDir,debug=debug)
    input('运行结束，回车退出')


if __name__ == '__main__':
    main()